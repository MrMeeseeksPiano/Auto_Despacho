from Auto_Despacho_get_links import texto_combinado
from docx import Document
from docx.shared import RGBColor
import datetime
from Auto_Despacho_get_links import dia_escolhido, mes_escolhido, ano_escolhido

hoje = datetime.date.today()
dia = hoje.day
mes = hoje.month
ano = hoje.year

# String original
texto = texto_combinado
# texto_emergencia =
# texto = texto_emergencia

# Palavras ou frases que você deseja destacar (em minúsculas para comparação)
palavra_azul = "operação comercial"
palavra_amarela = "operação em teste"
palavra_verde = "alterar as características técnicas"

# Cria um novo documento Word
doc = Document()

# Adiciona um parágrafo
p = doc.add_paragraph()

# Converta o texto para minúsculas apenas para fazer a comparação
texto_lower = texto.lower()

# Posição inicial do processamento do texto
pos = 0

while pos < len(texto):
    # Encontra a próxima ocorrência de uma das palavras
    pos_azul = texto_lower.find(palavra_azul, pos)
    pos_amarela = texto_lower.find(palavra_amarela, pos)
    pos_verde = texto_lower.find(palavra_verde, pos)

    # Se não houver mais palavras a serem destacadas, adicione o resto do texto
    if pos_azul == -1 and pos_amarela == -1 and pos_verde == -1:
        p.add_run(texto[pos:])
        break

    # Determine qual palavra aparece primeiro
    proxima_palavra = min(
        [pos_azul if pos_azul != -1 else len(texto),
         pos_amarela if pos_amarela != -1 else len(texto),
         pos_verde if pos_verde != -1 else len(texto)]
    )

    if proxima_palavra == pos_azul:
        # Adiciona o texto até a palavra "operação comercial"
        p.add_run(texto[pos:pos_azul])

        # Adiciona a palavra "operação comercial" em azul
        run_azul = p.add_run(texto[pos_azul:pos_azul+len(palavra_azul)])
        run_azul.bold = True
        run_azul.font.color.rgb = RGBColor(0, 0, 255)  # Cor azul (RGB)

        # Avança a posição
        pos = pos_azul + len(palavra_azul)

    elif proxima_palavra == pos_amarela:
        # Adiciona o texto até a palavra "operação em teste"
        p.add_run(texto[pos:pos_amarela])

        # Adiciona a palavra "operação em teste" em amarelo
        run_amarelo = p.add_run(
            texto[pos_amarela:pos_amarela+len(palavra_amarela)])
        run_amarelo.bold = True
        run_amarelo.font.color.rgb = RGBColor(255, 140, 0)  # Cor Laranja (RGB)

        # Avança a posição
        pos = pos_amarela + len(palavra_amarela)

    elif proxima_palavra == pos_verde:
        # Adiciona o texto até a palavra "alterar as características técnicas"
        p.add_run(texto[pos:pos_verde])

        # Adiciona a palavra "alterar as características técnicas" em verde
        run_verde = p.add_run(texto[pos_verde:pos_verde+len(palavra_verde)])
        run_verde.bold = True
        run_verde.font.color.rgb = RGBColor(0, 255, 0)  # Cor verde (RGB)

        # Avança a posição
        pos = pos_verde + len(palavra_verde)

# Salva o documento
doc.save(
    f'Despachos_texto_com_Link_{dia_escolhido}_{mes_escolhido}_{ano_escolhido}.docx')

input("Pressione Enter para sair...")
